"""
Unified format service for transcription results
Handles conversion to various output formats including text, SRT, VTT, PDF, and DOCX
"""
import io
import base64
import json
from datetime import timedelta
from typing import Dict, Any, List, Optional
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches

try:
    # Try relative import first (for FastAPI app)
    from ..utils.logger import get_logger
except ImportError:
    # Fall back to absolute import (for RQ worker script)
    from utils.logger import get_logger

class TranscriptionFormatService:
    """Unified service for converting transcription results to various output formats"""
    
    def __init__(self):
        self.logger = get_logger("format_service")
        self.supported_formats = ["text", "json", "srt", "vtt", "pdf", "docx"]
    
    async def convert_to_format(self, raw_data: Dict[str, Any], format_type: str, task_id: str = None) -> Dict[str, Any]:
        """
        Convert raw transcription data to the specified format
        
        Args:
            raw_data: Raw transcription data with segments, text, speakers, etc.
            format_type: Target format ('text', 'json', 'srt', 'vtt', 'pdf', 'docx')
            task_id: Optional task ID for file naming in PDF/DOCX
            
        Returns:
            Dictionary with the converted data
        """
        if not self.validate_format(format_type):
            raise ValueError(f"Unsupported format: {format_type}. Supported: {self.supported_formats}")
        
        try:
            if format_type == "text":
                return {"text": raw_data.get('text', '')}
            
            elif format_type == "json":
                return raw_data
            
            elif format_type == "srt":
                srt_content = self._generate_srt(raw_data.get('segments', []))
                return {"srt": srt_content}
            
            elif format_type == "vtt":
                vtt_content = self._generate_vtt(raw_data.get('segments', []))
                return {"vtt": vtt_content}
            
            elif format_type == "pdf":
                pdf_content = await self._generate_pdf(raw_data, task_id or "transcription")
                pdf_base64 = base64.b64encode(pdf_content).decode('utf-8')
                return {"pdf_base64": pdf_base64}
            
            elif format_type == "docx":
                docx_content = await self._generate_docx(raw_data, task_id or "transcription")
                docx_base64 = base64.b64encode(docx_content).decode('utf-8')
                return {"docx_base64": docx_base64}
            
        except Exception as e:
            self.logger.error(f"Failed to convert to {format_type}: {e}")
            raise Exception(f"Format conversion failed: {e}")
    
    def _generate_srt(self, segments: List[Dict[str, Any]]) -> str:
        """Generate SRT format from segments"""
        srt_content = []
        
        for i, segment in enumerate(segments, 1):
            start_time = self._format_time_srt(segment.get('start', 0))
            end_time = self._format_time_srt(segment.get('end', 0))
            text = segment.get('text', '').strip()
            speaker = segment.get('speaker', '')
            
            # Add speaker prefix if available
            if speaker and speaker != "SPEAKER_UNKNOWN":
                text = f"[{speaker}] {text}"
            
            srt_content.append(f"{i}")
            srt_content.append(f"{start_time} --> {end_time}")
            srt_content.append(text)
            srt_content.append("")  # Empty line between segments
        
        return "\n".join(srt_content)
    
    def _generate_vtt(self, segments: List[Dict[str, Any]]) -> str:
        """Generate VTT format from segments"""
        vtt_content = ["WEBVTT", ""]
        
        for segment in segments:
            start_time = self._format_time_vtt(segment.get('start', 0))
            end_time = self._format_time_vtt(segment.get('end', 0))
            text = segment.get('text', '').strip()
            speaker = segment.get('speaker', '')
            
            # Use VTT speaker notation
            if speaker and speaker != "SPEAKER_UNKNOWN":
                text = f"<v {speaker}>{text}"
            
            vtt_content.append(f"{start_time} --> {end_time}")
            vtt_content.append(text)
            vtt_content.append("")  # Empty line between segments
        
        return "\n".join(vtt_content)
    
    async def _generate_pdf(self, result: Dict[str, Any], task_id: str) -> bytes:
        """Generate PDF from transcription result"""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor='black',
                alignment=1  # Center alignment
            )
            
            story.append(Paragraph("Audio Transcription Report", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Metadata
            metadata_style = styles['Normal']
            story.append(Paragraph(f"<b>Task ID:</b> {task_id}", metadata_style))
            story.append(Paragraph(f"<b>Language:</b> {result.get('language', 'Unknown')}", metadata_style))
            story.append(Paragraph(f"<b>Duration:</b> {result.get('duration', 0):.2f} seconds", metadata_style))
            story.append(Paragraph(f"<b>Word Count:</b> {result.get('word_count', 0)}", metadata_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Full transcript
            story.append(Paragraph("<b>Full Transcript:</b>", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            full_text = result.get('text', '')
            if full_text:
                # Handle long text by splitting into paragraphs
                paragraphs = full_text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), styles['Normal']))
                        story.append(Spacer(1, 0.1*inch))
            
            story.append(Spacer(1, 0.3*inch))
            
            # Speaker segments (if available)
            segments = result.get('segments', [])
            if segments and any('speaker' in seg for seg in segments):
                story.append(Paragraph("<b>Speaker Segments:</b>", styles['Heading2']))
                story.append(Spacer(1, 0.1*inch))
                
                for segment in segments:
                    start_time = segment.get('start', 0)
                    end_time = segment.get('end', 0)
                    speaker = segment.get('speaker', 'Unknown')
                    text = segment.get('text', '').strip()
                    
                    time_range = f"[{start_time:.1f}s - {end_time:.1f}s]"
                    speaker_text = f"<b>{speaker}</b> {time_range}: {text}"
                    
                    story.append(Paragraph(speaker_text, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            self.logger.error(f"PDF generation failed: {e}")
            raise Exception(f"PDF generation failed: {e}")
    
    async def _generate_docx(self, result: Dict[str, Any], task_id: str) -> bytes:
        """Generate DOCX from transcription result"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Audio Transcription Report', 0)
            title.alignment = 1  # Center alignment
            
            # Metadata
            doc.add_heading('Metadata', level=1)
            
            metadata_table = doc.add_table(rows=4, cols=2)
            metadata_table.style = 'Table Grid'
            
            metadata_table.cell(0, 0).text = 'Task ID'
            metadata_table.cell(0, 1).text = task_id
            
            metadata_table.cell(1, 0).text = 'Language'
            metadata_table.cell(1, 1).text = result.get('language', 'Unknown')
            
            metadata_table.cell(2, 0).text = 'Duration'
            metadata_table.cell(2, 1).text = f"{result.get('duration', 0):.2f} seconds"
            
            metadata_table.cell(3, 0).text = 'Word Count'
            metadata_table.cell(3, 1).text = str(result.get('word_count', 0))
            
            # Full transcript
            doc.add_heading('Full Transcript', level=1)
            full_text = result.get('text', '')
            if full_text:
                # Handle long text by splitting into paragraphs
                paragraphs = full_text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            
            # Speaker segments (if available)
            segments = result.get('segments', [])
            if segments and any('speaker' in seg for seg in segments):
                doc.add_heading('Speaker Segments', level=1)
                
                for segment in segments:
                    start_time = segment.get('start', 0)
                    end_time = segment.get('end', 0)
                    speaker = segment.get('speaker', 'Unknown')
                    text = segment.get('text', '').strip()
                    
                    time_range = f"[{start_time:.1f}s - {end_time:.1f}s]"
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{speaker} ").bold = True
                    p.add_run(f"{time_range}: ")
                    p.add_run(text)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            self.logger.error(f"DOCX generation failed: {e}")
            raise Exception(f"DOCX generation failed: {e}")
    
    def _format_time_srt(self, seconds: float) -> str:
        """Format time for SRT (HH:MM:SS,mmm)"""
        td = timedelta(seconds=seconds)
        hours, remainder = divmod(td.total_seconds(), 3600)
        minutes, seconds = divmod(remainder, 60)
        milliseconds = int((seconds % 1) * 1000)
        
        return f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d},{milliseconds:03d}"
    
    def _format_time_vtt(self, seconds: float) -> str:
        """Format time for VTT (HH:MM:SS.mmm)"""
        td = timedelta(seconds=seconds)
        hours, remainder = divmod(td.total_seconds(), 3600)
        minutes, seconds = divmod(remainder, 60)
        milliseconds = int((seconds % 1) * 1000)
        
        return f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d}.{milliseconds:03d}"
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported output formats"""
        return self.supported_formats.copy()
    
    def validate_format(self, format_type: str) -> bool:
        """Validate if the format is supported"""
        return format_type in self.supported_formats
    
    def get_content_type(self, format_type: str) -> str:
        """Get appropriate content type for format"""
        content_types = {
            "text": "text/plain",
            "json": "application/json",
            "srt": "text/srt", 
            "vtt": "text/vtt",
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
        return content_types.get(format_type, "text/plain")
    
    def get_file_extension(self, format_type: str) -> str:
        """Get appropriate file extension for format"""
        extensions = {
            "text": ".txt",
            "json": ".json", 
            "srt": ".srt",
            "vtt": ".vtt",
            "pdf": ".pdf",
            "docx": ".docx"
        }
        return extensions.get(format_type, ".txt")

# Global instance
format_service = TranscriptionFormatService()